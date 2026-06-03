import streamlit as st
import docx
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re


def force_heading_bold(paragraph):
    """
    Forza il bold a livello di rPr nell'XML dei run del paragrafo,
    necessario perché i temi di Word sovrascrivono altrimenti il run.
    """
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        if not rPr.xpath('w:b'):
            b = OxmlElement('w:b')
            rPr.append(b)
        if not rPr.xpath('w:bCs'):
            b_cs = OxmlElement('w:bCs')
            rPr.append(b_cs)


def set_run_font(run, name, size_pt, bold=False, color_rgb=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    # Forza anche il font nei tag XML del run per evitare l'override del tema
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.insert(0, rFonts)


def parse_inline_markdown(p, text, base_font='Garamond', base_size=12):
    """
    Parsea markdown inline (**grassetto** e *corsivo*) e aggiunge run al paragrafo.
    """
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, base_font, base_size, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_run_font(run, base_font, base_size)
            run.font.italic = True
        else:
            run = p.add_run(part)
            set_run_font(run, base_font, base_size)


def add_page_number(run):
    """
    Aggiunge il campo dinamico PAGE (numero di pagina) nell'XML del run.
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def add_toc(paragraph):
    """
    Aggiunge un Indice (Table of Contents) aggiornabile in Word.
    """
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def format_thesis_doc(testo, university):
    doc = docx.Document()
    
    # Imposta i parametri stilistici a seconda dell'università
    if university == "IULM":
        font_name = 'Garamond'
        line_spacing = 1.5
        margins = (3.0, 3.0, 3.0, 3.0)  # top, bottom, left, right in cm
        empty_pages = 2
        page_num_align = WD_ALIGN_PARAGRAPH.LEFT
    elif university == "Bocconi":
        font_name = 'Arial'
        line_spacing = 1.15  # interlinea per contenere 26-30 righe per pagina
        margins = (2.5, 2.5, 2.5, 2.5)
        empty_pages = 4
        page_num_align = WD_ALIGN_PARAGRAPH.RIGHT
    elif university == "Cattolica":
        font_name = 'Times New Roman'
        line_spacing = 1.5
        margins = (3.5, 2.5, 3.5, 2.5)  # asimmetria per rilegatura
        empty_pages = 0
        page_num_align = WD_ALIGN_PARAGRAPH.CENTER
    elif university == "Politecnico di Milano":
        font_name = 'Times New Roman'
        line_spacing = 1.5
        margins = (3.0, 3.0, 3.0, 2.5)
        empty_pages = 0
        page_num_align = WD_ALIGN_PARAGRAPH.CENTER
    elif university == "Università di Pavia":
        font_name = 'Times New Roman'
        line_spacing = 1.5
        margins = (3.0, 3.0, 3.0, 3.0)
        empty_pages = 0
        page_num_align = WD_ALIGN_PARAGRAPH.CENTER
    else:  # Generale (Standard)
        font_name = 'Times New Roman'
        line_spacing = 1.5
        margins = (3.0, 3.0, 3.5, 3.0)  # Margine sinistro di 3.5 per rilegatura
        empty_pages = 0
        page_num_align = WD_ALIGN_PARAGRAPH.CENTER

    # --- Sezione 1: Frontespizio ---
    sect_frontespizio = doc.sections[0]
    sect_frontespizio.top_margin = Cm(margins[0])
    sect_frontespizio.bottom_margin = Cm(margins[1])
    sect_frontespizio.left_margin = Cm(margins[2])
    sect_frontespizio.right_margin = Cm(margins[3])
    sect_frontespizio.footer.is_linked_to_previous = False

    # Stile Normal
    normal = doc.styles['Normal']
    normal.font.name = font_name
    normal.font.size = Pt(12)
    normal.paragraph_format.widow_control = True

    # Stile Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = font_name
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)

    # Stile Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = font_name
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(6)

    # Inserisce pagine vuote se richiesto
    for _ in range(empty_pages):
        doc.add_page_break()

    # --- Inserimento Indice (TOC) ---
    p_toc_title = doc.add_heading('', level=1)
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_toc = p_toc_title.add_run('Indice')
    set_run_font(run_toc, font_name, 20, bold=True, color_rgb=RGBColor(0, 0, 0))
    force_heading_bold(p_toc_title)
    
    p_toc = doc.add_paragraph()
    add_toc(p_toc)

    p_toc_inst = doc.add_paragraph()
    run_inst = p_toc_inst.add_run(" (Dopo aver aperto il file Word, clicca col tasto destro in questo spazio vuoto e seleziona 'Aggiorna campo' -> 'Aggiorna intero sommario' per far apparire l'indice corretto. Devi farlo a mano altrimenti Word sbaglia i numeri!)")
    set_run_font(run_inst, font_name, 10, bold=False, color_rgb=RGBColor(128, 128, 128))

    first_chapter = True
    in_bibliography = False
    
    lines = testo.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # --- Sostituzioni Matematiche di Base ---
        line = line.replace('$R^2$', 'R²')
        line = line.replace('R^2', 'R²')
        line = line.replace('beta', 'β')
        line = line.replace('alpha', 'α')
        line = line.replace('gamma', 'γ')
        line = line.replace('epsilon', 'ε')
        line = line.replace('lambda', 'λ')

        # --- Tabella Markdown ---
        if line.startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            valid_rows = []
            for t_line in table_lines:
                if re.match(r'^\|[\s\-\|:]+\|$', t_line):
                    continue
                cells = [c.strip() for c in t_line.split('|')[1:-1]]
                valid_rows.append(cells)

            if valid_rows:
                num_cols = len(valid_rows[0])
                table = doc.add_table(rows=len(valid_rows), cols=num_cols)
                table.style = 'Table Grid'
                
                # Forza l'autofit al 100% della pagina
                tblPr = table._tbl.tblPr
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:type'), 'pct')
                tblW.set(qn('w:w'), '5000') # 5000 = 100%
                tblPr.append(tblW)

                for row_idx, row_data in enumerate(valid_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell_clean = cell_text.replace('**', '').replace('*', '')
                            cell.text = cell_clean
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.font.name = font_name
                                    run.font.size = Pt(11)
                                    if row_idx == 0:
                                        run.font.bold = True
            continue

        # --- Heading 2 (##) ---
        if line.startswith('## '):
            title_text = line[3:]
            p = doc.add_heading('', level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(title_text)
            set_run_font(run, font_name, 16, bold=True, color_rgb=RGBColor(0, 0, 0))
            force_heading_bold(p)

        # --- Heading 1 (#) ---
        elif line.startswith('# '):
            title_text = line[2:]
            
            if title_text.lower() == 'bibliografia':
                in_bibliography = True
                
            if first_chapter:
                # Se è il primo capitolo, creiamo la Sezione 2 con la formattazione dell'ateneo
                new_section = doc.add_section(WD_SECTION_START.ODD_PAGE)
                new_section.top_margin = Cm(margins[0])
                new_section.bottom_margin = Cm(margins[1])
                new_section.left_margin = Cm(margins[2])
                new_section.right_margin = Cm(margins[3])
                
                # Sgancia il footer
                new_section.footer.is_linked_to_previous = False
                
                # Reset numerazione a 1 (tranne Bocconi che vuole 4 pagine bianche e ha regole particolari)
                sectPr = new_section._sectPr
                pgNumType = OxmlElement('w:pgNumType')
                pgNumType.set(qn('w:start'), "1")
                sectPr.append(pgNumType)
                
                # Aggiungi il numero pagina nel footer (se non è Bocconi che preferisce anonimato totale)
                footer = new_section.footer
                if len(footer.paragraphs) == 0:
                    p_footer = footer.add_paragraph()
                else:
                    p_footer = footer.paragraphs[0]
                    
                p_footer.alignment = page_num_align
                run_footer = p_footer.add_run()
                run_footer.font.name = font_name
                run_footer.font.size = Pt(12)
                add_page_number(run_footer)
                
                first_chapter = False
            else:
                new_section = doc.add_section(WD_SECTION_START.ODD_PAGE)
                new_section.top_margin = Cm(margins[0])
                new_section.bottom_margin = Cm(margins[1])
                new_section.left_margin = Cm(margins[2])
                new_section.right_margin = Cm(margins[3])
                
                # Continua la numerazione
                sectPr = new_section._sectPr
                pgNumType = sectPr.find(qn('w:pgNumType'))
                if pgNumType is not None:
                    sectPr.remove(pgNumType)
                    
            first_chapter = False

            p = doc.add_heading('', level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            set_run_font(run, font_name, 20, bold=True, color_rgb=RGBColor(0, 0, 0))
            force_heading_bold(p)

        # --- Testo normale ---
        else:
            p = doc.add_paragraph()
            parse_inline_markdown(p, line, base_font=font_name)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = line_spacing
            if in_bibliography:
                p.paragraph_format.space_after = Pt(12)
            else:
                p.paragraph_format.space_after = Pt(0)

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- UI Streamlit ---
st.set_page_config(page_title="FormatTesi.it — Multi-University Formatter", layout="wide", initial_sidebar_state="expanded")

st.title("FormatTesi.it")
st.markdown(
    "Un ambiente professionale per formattare la tua tesi di laurea in conformità alle linee guida ufficiali degli atenei italiani."
)

# Parsing per Indice
if "testo_input" not in st.session_state:
    st.session_state.testo_input = ""

def update_text():
    st.session_state.testo_input = st.session_state.editor_area

with st.sidebar:
    st.header("Opzioni Template")
    univ_selected = st.selectbox(
        "Seleziona il template dell'ateneo:",
        ["Generale (Standard)", "IULM", "Bocconi", "Cattolica", "Politecnico di Milano", "Università di Pavia"]
    )
    
    st.markdown("---")
    st.header("Struttura Tesi")
    if st.session_state.testo_input.strip():
        lines = st.session_state.testo_input.split('\n')
        h_counter = 0
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('# '):
                title = stripped[2:]
                h_counter += 1
                anchor = f"capitolo-{h_counter}"
                st.markdown(f"**<a href='#{anchor}' style='color: inherit; text-decoration: none;'>{title}</a>**", unsafe_allow_html=True)
            elif stripped.startswith('## '):
                title = stripped[3:]
                h_counter += 1
                anchor = f"paragrafo-{h_counter}"
                st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;<a href='#{anchor}' style='color: inherit; text-decoration: none;'>{title}</a>", unsafe_allow_html=True)
    else:
        st.info("Incolla il testo per generare l'indice navigabile.")

# Layout a Tab
tab1, tab2 = st.tabs(["Editor Testuale", "Anteprima di Lettura"])

with tab1:
    st.markdown("### Area di Inserimento")
    st.text_area(
        "Incolla qui il tuo testo in markdown (usa # per i Capitoli e ## per i Paragrafi):", 
        height=600, 
        key="editor_area", 
        on_change=update_text
    )

    st.markdown("---")
    colA, colB = st.columns([1, 3])
    with colA:
        if st.button("Formatta e genera Word", type="primary", use_container_width=True):
            if st.session_state.testo_input.strip():
                with st.spinner("Generazione del documento Word in corso..."):
                    docx_file = format_thesis_doc(st.session_state.testo_input, univ_selected)
                st.success("Generazione completata!")
                
                clean_univ_name = univ_selected.replace(' ', '_').replace('(', '').replace(')', '')
                st.download_button(
                    label="Scarica Tesi (.docx)",
                    data=docx_file,
                    file_name=f"Tesi_{clean_univ_name}_Formattata.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.error("Inserisci del testo prima di generare!")

with tab2:
    st.markdown("### Preview Navigabile")
    if st.session_state.testo_input.strip():
        with st.container(height=700):
            lines = st.session_state.testo_input.split('\n')
            out_lines = []
            h_counter = 0
            
            for line in lines:
                stripped = line.strip()
                if stripped.startswith('# '):
                    title = stripped[2:]
                    h_counter += 1
                    anchor = f"capitolo-{h_counter}"
                    out_lines.append(f"<a id='{anchor}'></a>\n\n# {title}")
                elif stripped.startswith('## '):
                    title = stripped[3:]
                    h_counter += 1
                    anchor = f"paragrafo-{h_counter}"
                    out_lines.append(f"<a id='{anchor}'></a>\n\n## {title}")
                else:
                    out_lines.append(line)
            
            st.markdown("\n".join(out_lines), unsafe_allow_html=True)
    else:
        st.info("L'anteprima apparirà qui non appena avrai incollato il testo nella tab Editor.")
